import streamlit as st
from io import BytesIO
from docx import Document
from datetime import datetime

st.set_page_config(page_title="Therapy Feedback Form", layout="centered")

# -------------------------------
# Helper: record full question text + response
# -------------------------------
def ask(question, widget_func, *args, **kwargs):
    """Renders a Streamlit widget and automatically stores question/answer."""
    answer = widget_func(question, *args, **kwargs)
    st.session_state["_current_section"][question] = answer
    return answer

# -------------------------------
# Helper: Generate Word doc
# -------------------------------
def make_docx(responses, title="Therapy Feedback Summary"):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%SZ')} (UTC)\n")

    doc.add_paragraph(
        "Summary for Therapists:\n"
        "This document compiles client feedback from the therapy reflection tool. "
        "All responses are provided as entered by the client. "
        "Only answered questions are included. Likert values (1–5) reflect selected ratings. "
        "Blank answers are excluded."
    )
    doc.add_paragraph()

    for i, (section, qas) in enumerate(responses.items(), start=1):
        answered = {q: a for q, a in qas.items() if a not in [None, "", [], [""]]}
        if not answered:
            continue
        doc.add_heading(f"Section {i}: {section}", level=2)
        for q, a in answered.items():
            p = doc.add_paragraph()
            p.add_run(f"{q}: ").bold = True
            p.add_run(str(a))
        doc.add_paragraph()

    doc.add_paragraph("Note: Only answered questions and selected sections are included.")
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f.read()

# -------------------------------
# Initialize session state
# -------------------------------
if "responses" not in st.session_state:
    st.session_state["responses"] = {}

# -------------------------------
# Introduction
# -------------------------------
st.title("Therapy Feedback Form")
st.warning("⚠️ This is an experimental / early version of the tool We would recommend completing in one sitting if possible and checking the final summary to ensure it looks the way you want.")
st.markdown("""
This form helps you reflect on your therapy experiences — what’s working, what feels hard, and what you might want to share. 
All questions are optional and only answered questions will appear in the final download. 
You can answer as many or as few as you like, and save everything at the end using the **“Save All Sections”** button.
""")
st.markdown("---")

# -------------------------------
# Section 0 — Entry Point
# -------------------------------
st.header("Section 0 — Entry Point")

st.session_state["_current_section"] = {}
focus = ask(
    "Which areas would you like this form to focus on today?",
    st.radio,
    options=("General Feedback", "Concerns, Discomfort, or Harm", "Both"),
    key="entry_focus"
)
style = ask(
    "Preferred style of questions/prompts:",
    st.radio,
    options=["Structured", "Unstructured"],
    format_func=lambda s: "Includes Likert scales and structured prompts"
        if s == "Structured" else "Focuses on short open-ended reflection",
    key="entry_style"
)
harm_expanded = focus in ["Concerns, Discomfort, or Harm", "Both"]
st.session_state["responses"]["Entry Point"] = st.session_state["_current_section"]

st.markdown("---")

# -------------------------------
# Section 1 — Overall Experience
# -------------------------------
st.header("Section 1 — Overall Experience")
st.session_state["_current_section"] = {}

if focus in ["General Feedback", "Both"]:
    if style == "Structured":
        ask("Overall, how well is therapy meeting your needs? (1–5)", st.selectbox, ["", 1, 2, 3, 4, 5], index=0)
        ask("Anything you especially enjoy or hope continues?", st.text_area, height=100)
        ask("Anything stressful, uncomfortable, or less effective?", st.text_area, height=100)
    else:
        ask("What is most helpful about therapy?", st.text_area, height=100)
        ask("What feels difficult or could be improved?", st.text_area, height=100)
    ask("What do you experience me as least aware of in how I’m engaging with you? Were I to become more aware of it, what do you imagine would change?", st.text_area)
    ask("Additional comments / clarifications for this section", st.text_area, key="additional_comments_oe")

st.session_state["responses"]["Overall Experience"] = st.session_state["_current_section"]

# -------------------------------
# Section 2 — Topics & Comfort
# -------------------------------
with st.expander("Section 2 — Topics & Comfort"):
    st.session_state["_current_section"] = {}
    if style == "Structured":
        ask("What topics are easy to discuss with your therapist?", st.text_area)
        ask("What topics are more challenging to discuss?", st.text_area)
        ask("Any topics you’ve been hesitant to bring up?", st.text_area)
        ask("Trainings or resources you wish your therapist knew about?", st.text_area)
    else:
        ask("Topics you feel comfortable discussing?", st.text_area)
        ask("Topics that feel difficult?", st.text_area)
    ask("Additional comments / clarifications for this section", st.text_area, key="additional_comments_tc")
    st.session_state["responses"]["Topics & Comfort"] = st.session_state["_current_section"]

# -------------------------------
# Section 3 — Structure & Style
# -------------------------------
with st.expander("Section 3 — Structure & Style"):
    st.session_state["_current_section"] = {}
    if style == "Structured":
        ask("Punctuality/scheduling is a source of stress (1–5)", st.selectbox, ["", 1, 2, 3, 4, 5], index=0)
        ask("Therapist self-disclosure (i.e. how much the therapist shares about themself)", st.selectbox, ["", "Would prefer less self-disclosure", "The current amount is good", "Would prefer more self-disclosure", "Unsure/conflicting feelings"], index=0)
        ask("Is there anything else you would like to add about self-disclosure?", st.text_area)
        ask("Session structure/flexibility", st.selectbox, ["", "Would prefer more flexibility", "Balance is good", "Would prefer more structure", "Uncertain/conflicted"], index=0)
        ask("Is there anything else you would like to add about session structure/flexibility?", st.text_area)
        ask("Do you feel like your communication style is well-understood/supported?", st.text_area)
        ask("How well does therapy support your autonomy/ability to be yourself?", st.selectbox, ["", 1, 2, 3, 4, 5], index=0)
        ask("Thoughts on therapist availability or session frequency?", st.text_area)
        ask("Areas you wish therapist were more or less directive/proactive?", st.text_area)
    else:
        ask("How do you experience the structure of sessions?", st.text_area)
        ask("Anything you’d like to change about structure or therapist style?", st.text_area)
    ask("Is there anything you would like to change about the physical environment of therapy (e.g., seating arrangement, proximity to the door, level of lighting)?", st.text_area)
    ask("Additional comments / clarifications for this section", st.text_area, key="additional_comments_ss")
    st.session_state["responses"]["Structure & Style"] = st.session_state["_current_section"]


# -------------------------------
# Section 4 — Relational Climate
# -------------------------------
with st.expander("Section 4 — Relational Climate"):
    st.session_state["_current_section"] = {}
    if style == "Structured":
        ask("I feel safe, respected, and understood (1–5)", st.selectbox, ["", 1, 2, 3, 4, 5], index=0)
        ask("I feel comfortable being honest (1–5)", st.selectbox, ["", 1, 2, 3, 4, 5], index=0)
        ask("Overall emotional experiences after sessions (check all that apply)", st.multiselect, ["Calm","Relieved","Overwhelmed","Confused","Hopeful","Worse"], default=[])
        ask("Signs you notice if you start to feel stressed during a session", st.text_area)
        ask("What helps most when stressed during or after a session?", st.text_area)
    else:
        ask("Emotional experiences after sessions?", st.text_area)
        ask("What helps when stressed during/after sessions?", st.text_area)
    ask("Additional comments / clarifications for this section", st.text_area, key="additional_comments_rc")
    st.session_state["responses"]["Relational Climate"] = st.session_state["_current_section"]

# -------------------------------
# Section 5 — Therapy Harm & Boundaries
# -------------------------------
with st.expander("Section 5 — Therapy Harm & Boundaries", expanded=harm_expanded):
    st.session_state["_current_section"] = {}
    if style == "Structured":
        ask("I sometimes feel like my therapist treats me more like a friend than a client (1–5)", st.selectbox, ["",1,2,3,4,5], index=0)
        ask("I sometimes feel uncomfortable with how much my therapist shares (1–5)", st.selectbox, ["",1,2,3,4,5], index=0)
        ask("I feel confused or hurt about physical contact (1–5)", st.selectbox, ["",1,2,3,4,5], index=0)
        ask("Dual roles: overlap in social/professional contexts?", st.text_area)
        ask("Any confusing moments regarding boundaries that affected you?", st.text_area)
        ask("I sometimes feel dismissed or unheard (1–5)", st.selectbox, ["",1,2,3,4,5], index=0)
        ask("Impact of these experiences", st.text_area)
        ask("What acknowledgement or response would feel meaningful?", st.text_area)
        ask("Any financial concerns (e.g., overbilling, pressured sessions)?", st.text_area)
        ask("Moments of intense emotional highs/lows or confusing attachment feelings?", st.text_area)
        ask("Insensitive behavior regarding identity or experience?", st.text_area)
        ask("Anything else about harm, safety, or boundaries?", st.text_area)
    else:
        ask("Any experiences of concern, harm, or boundary issues you want to note?", st.text_area)
    ask("Additional comments / clarifications for this section", st.text_area, key="additional_comments_thb")
    st.session_state["responses"]["Therapy Harm & Boundaries"] = st.session_state["_current_section"]

# -------------------------------
# Section 6 — Feedback & Wrap-Up
# -------------------------------
with st.expander("Section 6 — Feedback & Wrap-Up"):
    st.session_state["_current_section"] = {}
    if style == "Structured":
        ask("How would you ideally like your therapist to use this feedback?", st.multiselect,
            ["Discuss in session", "Read privately", "Acknowledge through email", "Share how they plan to make changes", "Other"], default=[])
        ask("Emotions when sharing feedback (select all that apply)", st.multiselect,
            ["Comfortable","Nervous","Afraid","Uncertain","Relieved","Secure","Conflicted","Feeling guilt"], default=[])
        ask("Any fears about sharing feedback (select all that apply)", st.multiselect,
            ["Receiving a defensive response","Loss of connection or regard","Being forced into other treatment","Therapist ends therapy","Other"], default=[])
        ask("Additional thoughts or clarifications about sharing feedback / supports", st.text_area)
    else:
        ask("Any thoughts about sharing feedback or support needs?", st.text_area)
    st.session_state["responses"]["Feedback & Wrap-Up"] = st.session_state["_current_section"]

# -------------------------------
# Save and Download
# -------------------------------
st.markdown("---")
st.header("Create and Download Summary")

include_sections = st.multiselect(
    "Select sections to include:",
    list(st.session_state["responses"].keys()),
    default=list(st.session_state["responses"].keys())
)
filtered = {sec: st.session_state["responses"][sec] for sec in include_sections if st.session_state["responses"][sec]}

if st.button("Generate Word Summary"):
    if not filtered:
        st.warning("No responses yet! Please fill and save sections first.")
    else:
        docx_bytes = make_docx(filtered)
        st.session_state["docx_bytes"] = docx_bytes
        st.success("Summary generated!")

if "docx_bytes" in st.session_state:
    filename = f"therapy_feedback_{datetime.utcnow().strftime('%Y%m%d_%H%M')}.docx"
    st.download_button(
        "⬇️ Download Word Summary",
        data=st.session_state["docx_bytes"],
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


